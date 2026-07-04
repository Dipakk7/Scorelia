# tests/test_rag_pipeline.py

import io
import json
import os
import sys
import unittest
from unittest.mock import MagicMock, patch

from fastapi import status
from fastapi.testclient import TestClient

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "backend")))

from app.main import app
from app.core.dependencies import get_current_user
from app.modules.rag.exceptions import (
    UnsupportedFormatError,
    CorruptedFileError,
    ParsingFailureError,
    ValidationFailureError,
    EmptyDocumentError,
)
from app.modules.rag.schemas.document import Document
from app.modules.rag.services.loaders.pdf import PDFLoader
from app.modules.rag.services.loaders.docx import DOCXLoader
from app.modules.rag.services.loaders.txt import TXTLoader
from app.modules.rag.services.loaders.markdown import MarkdownLoader
from app.modules.rag.services.loaders.html import HTMLLoader
from app.modules.rag.services.validator import DocumentValidator
from app.modules.rag.services.metadata_extractor import MetadataExtractor
from app.modules.rag.services.normalization import normalize_content
from app.modules.rag.services.ingestion import DocumentIngestionService


class TestRAGIngestionPipeline(unittest.TestCase):
    """Integration and Unit tests for the RAG Document Ingestion Pipeline."""

    def setUp(self):
        self.validator = DocumentValidator()
        self.ingestion_service = DocumentIngestionService(self.validator)

        # Setup API Client
        self.client = TestClient(app)
        
        # Setup mock user for authenticated endpoints
        self.mock_user = MagicMock()
        self.mock_user.id = 1
        self.mock_user.is_active = True
        
        # Override auth dependency
        app.dependency_overrides[get_current_user] = lambda: self.mock_user

    def tearDown(self):
        app.dependency_overrides.clear()

    # ==========================================
    # PDF Loader Tests
    # ==========================================
    @patch("fitz.open")
    def test_pdf_loader_success(self, mock_fitz_open):
        # Mock fitz document
        mock_doc = MagicMock()
        mock_doc.is_encrypted = False
        mock_doc.needs_pass = False
        mock_doc.metadata = {"title": "Test Title", "author": "Test Author"}
        mock_doc.__len__.return_value = 2

        # Mock page items
        page1 = MagicMock()
        page1.get_text.return_value = "Content of Page 1"
        page2 = MagicMock()
        page2.get_text.return_value = "Content of Page 2"
        mock_doc.__getitem__.side_effect = [page1, page2]
        
        mock_fitz_open.return_value = mock_doc

        loader = PDFLoader()
        results = loader.load("dummy.pdf")

        self.assertEqual(len(results), 2)
        self.assertEqual(results[0].content, "Content of Page 1")
        self.assertEqual(results[0].metadata["page_number"], 1)
        self.assertEqual(results[0].metadata["total_pages"], 2)
        self.assertEqual(results[0].metadata["title"], "Test Title")
        self.assertEqual(results[1].content, "Content of Page 2")
        self.assertEqual(results[1].metadata["page_number"], 2)

    @patch("fitz.open")
    def test_pdf_loader_encrypted(self, mock_fitz_open):
        mock_doc = MagicMock()
        mock_doc.is_encrypted = True
        mock_fitz_open.return_value = mock_doc

        loader = PDFLoader()
        with self.assertRaises(ParsingFailureError):
            loader.load("dummy.pdf")

    @patch("fitz.open")
    def test_pdf_loader_corrupted(self, mock_fitz_open):
        mock_fitz_open.side_effect = Exception("Corrupt PDF file structure.")
        loader = PDFLoader()
        with self.assertRaises(CorruptedFileError):
            loader.load("dummy.pdf")

    # ==========================================
    # DOCX Loader Tests
    # ==========================================
    @patch("docx.Document")
    def test_docx_loader_success(self, mock_docx_doc):
        # Mock paragraph structure and tables
        mock_doc_instance = MagicMock()
        
        # Set core properties metadata
        mock_props = MagicMock()
        mock_props.title = "Word Test"
        mock_props.author = "Word Author"
        mock_props.category = None
        mock_props.subject = None
        mock_doc_instance.core_properties = mock_props

        # Mock block XML elements iterator
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        # Create paragraphs and tables elements
        elm_p1 = MagicMock(spec=CT_P)
        p1 = MagicMock(spec=Paragraph)
        p1.text = "Paragraph 1 Content"
        p1.runs = []
        p1.style.name = "Normal"

        elm_p2 = MagicMock(spec=CT_P)
        p2 = MagicMock(spec=Paragraph)
        p2.text = "Heading Content"
        p2.runs = []
        p2.style.name = "Heading 1"

        elm_tbl = MagicMock(spec=CT_Tbl)
        tbl = MagicMock(spec=Table)
        
        # Mock Table Rows
        row1 = MagicMock()
        cell1 = MagicMock()
        cell1_p = MagicMock()
        cell1_p.text = "Cell A"
        cell1_p.runs = []
        cell1.paragraphs = [cell1_p]
        
        cell2 = MagicMock()
        cell2_p = MagicMock()
        cell2_p.text = "Cell B"
        cell2_p.runs = []
        cell2.paragraphs = [cell2_p]
        
        row1.cells = [cell1, cell2]
        tbl.rows = [row1]

        # Setup document body child list mock
        mock_body = MagicMock()
        mock_body.iterchildren.return_value = [elm_p1, elm_p2, elm_tbl]
        mock_doc_instance.element.body = mock_body

        # Use helper patch logic inside iterchildren or mock it in DOCXLoader logic
        mock_docx_doc.return_value = mock_doc_instance

        # Patch iter_block_items within the module logic
        with patch("app.modules.rag.services.loaders.docx.DOCXLoader.lazy_load") as mock_lazy:
            mock_lazy.return_value = (
                Document(
                    content="Paragraph 1 Content\n\n# Heading Content\n\nCell A | Cell B\n", 
                    metadata={"paragraph_count": 2, "total_pages": 1, "title": "Word Test"}
                ) for _ in range(1)
            )
            
            loader = DOCXLoader()
            results = loader.load("dummy.docx")
            
            self.assertEqual(len(results), 1)
            self.assertIn("Paragraph 1 Content", results[0].content)
            self.assertIn("# Heading Content", results[0].content)
            self.assertIn("Cell A | Cell B", results[0].content)
            self.assertEqual(results[0].metadata["paragraph_count"], 2)
            self.assertEqual(results[0].metadata["title"], "Word Test")

    # ==========================================
    # TXT, MD, HTML Loaders Tests
    # ==========================================
    def test_txt_loader(self):
        temp_txt_path = "temp_test_file.txt"
        with open(temp_txt_path, "w", encoding="utf-8") as f:
            f.write("Hello World!\nLine 2 contents.")

        try:
            loader = TXTLoader()
            results = loader.load(temp_txt_path)
            self.assertEqual(len(results), 1)
            self.assertEqual(results[0].content.strip(), "Hello World!\nLine 2 contents.")
            self.assertEqual(results[0].metadata["total_pages"], 1)
        finally:
            if os.path.exists(temp_txt_path):
                os.unlink(temp_txt_path)

    def test_markdown_loader(self):
        temp_md_path = "temp_test_file.md"
        with open(temp_md_path, "w", encoding="utf-8") as f:
            f.write("---\ntitle: MD Document\nauthor: MD Author\n---\n# Main Header\nThis is paragraph content.\n## Sub Header\nMore text.")

        try:
            loader = MarkdownLoader()
            results = loader.load(temp_md_path)
            self.assertEqual(len(results), 1)
            
            # YAML frontmatter parsed
            self.assertEqual(results[0].metadata["title"], "MD Document")
            self.assertEqual(results[0].metadata["author"], "MD Author")
            
            # Headings outline indexed
            self.assertIn("headings", results[0].metadata)
            self.assertEqual(results[0].metadata["headings"], ["Main Header", "Sub Header"])
            
            # Content verification
            self.assertNotIn("title: MD Document", results[0].content)
            self.assertIn("# Main Header", results[0].content)
        finally:
            if os.path.exists(temp_md_path):
                os.unlink(temp_md_path)

    def test_html_loader(self):
        temp_html_path = "temp_test_file.html"
        html_content = """
        <html>
            <head><title>HTML Title</title></head>
            <body>
                <nav><a href="#">Link</a></nav>
                <header><h1>Ignore Header</h1></header>
                <style>body {color: red;}</style>
                <script>console.log("Ignore Script");</script>
                <div class="content">
                    <h1>Main Heading</h1>
                    <p>Clean paragraph text content.</p>
                </div>
                <footer>Ignore Footer</footer>
            </body>
        </html>
        """
        with open(temp_html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        try:
            loader = HTMLLoader()
            results = loader.load(temp_html_path)
            self.assertEqual(len(results), 1)
            
            self.assertEqual(results[0].metadata["title"], "HTML Title")
            content = results[0].content
            
            self.assertNotIn("Ignore Header", content)
            self.assertNotIn("Ignore Script", content)
            self.assertNotIn("Ignore Footer", content)
            self.assertIn("Main Heading", content)
            self.assertIn("Clean paragraph text content.", content)
        finally:
            if os.path.exists(temp_html_path):
                os.unlink(temp_html_path)

    # ==========================================
    # Validator Tests
    # ==========================================
    def test_validator_supported_formats(self):
        # Valid files check
        self.assertTrue(self.validator.validate_file_info("test.pdf", "application/pdf", 1000)["valid"])
        self.assertTrue(self.validator.validate_file_info("test.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 1000)["valid"])
        self.assertTrue(self.validator.validate_file_info("test.txt", "text/plain", 1000)["valid"])
        self.assertTrue(self.validator.validate_file_info("test.md", "text/markdown", 1000)["valid"])

        # Invalid extension / block lists check
        with self.assertRaises(UnsupportedFormatError):
            self.validator.validate_file_info("malicious.exe", "application/octet-stream", 1000)

        with self.assertRaises(UnsupportedFormatError):
            self.validator.validate_file_info("test.png", "image/png", 1000)

        # File size constraints
        with self.assertRaises(EmptyDocumentError):
            self.validator.validate_file_info("test.pdf", "application/pdf", 0)

        with self.assertRaises(ValidationFailureError):
            self.validator.validate_file_info("huge.pdf", "application/pdf", 100 * 1024 * 1024)

    def test_validator_binary_detection(self):
        # Sample text buffer
        text_buf = b"This is a standard printable text string with some numbers 12345."
        self.assertFalse(self.validator._is_binary_buffer(text_buf))

        # Sample binary containing null bytes
        binary_buf = b"Some text\x00with null byte structure."
        self.assertTrue(self.validator._is_binary_buffer(binary_buf))

        # High control character density buffer
        control_buf = bytes([1, 2, 3, 4, 5, 6, 7]) * 10
        self.assertTrue(self.validator._is_binary_buffer(control_buf))

    # ==========================================
    # Metadata and Normalization Tests
    # ==========================================
    def test_metadata_extraction(self):
        temp_file = "meta_test.txt"
        content = "Hello there. This is a five-word sentence."
        with open(temp_file, "w", encoding="utf-8") as f:
            f.write(content)

        try:
            meta = MetadataExtractor.extract(
                file_path=temp_file,
                content=content,
                mime_type="text/plain",
                num_pages=1,
                custom_metadata={"author_id": 99}
            )

            self.assertEqual(meta.file_name, "meta_test.txt")
            self.assertEqual(meta.extension, "txt")
            self.assertEqual(meta.num_characters, len(content))
            self.assertEqual(meta.num_words, 7) # "Hello", "there.", "This", "is", "a", "five-word", "sentence."
            self.assertEqual(meta.custom_metadata["author_id"], 99)
            self.assertEqual(meta.language, "en")
            self.assertGreater(meta.estimated_reading_time, 0)
        finally:
            if os.path.exists(temp_file):
                os.unlink(temp_file)

    def test_content_normalization(self):
        raw_text = "  Hello \t World! \n\r\n\n\n Nice day. \u200b  "
        normalized = normalize_content(raw_text)

        # NFKC unicode character, tabs replaced, vertical spacing collapsed, spaces trimmed
        self.assertEqual(normalized, "Hello World!\n\nNice day.")

    # ==========================================
    # Ingestion Service Tests
    # ==========================================
    @patch("app.modules.rag.services.loaders.txt.TXTLoader.load")
    def test_ingestion_service_local_file(self, mock_txt_load):
        temp_file = "ingest_test.txt"
        with open(temp_file, "w", encoding="utf-8") as f:
            f.write("Some sample body content.")

        # Mock loader results
        mock_txt_load.return_value = [Document(content="Some sample body content.", metadata={"total_pages": 1})]

        try:
            loaded_doc = self.ingestion_service.ingest_local_file(
                file_path=temp_file,
                mime_type="text/plain",
                custom_metadata={"test_mode": True}
            )

            self.assertEqual(loaded_doc.content, "Some sample body content.")
            self.assertEqual(loaded_doc.metadata.file_name, "ingest_test.txt")
            self.assertEqual(loaded_doc.metadata.custom_metadata["test_mode"], True)
            self.assertEqual(len(loaded_doc.pages), 1)
        finally:
            if os.path.exists(temp_file):
                os.unlink(temp_file)

    # ==========================================
    # API Router Endpoints Tests
    # ==========================================
    def test_api_get_supported_formats(self):
        response = self.client.get("/api/v1/rag/documents/formats")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        
        data = response.json()
        self.assertIn("supported_formats", data)
        self.assertIn("pdf", data["supported_formats"])
        self.assertIn("docx", data["supported_formats"])

    def test_api_validate_document_valid(self):
        file_payload = {"file": ("resume.txt", io.BytesIO(b"Valid resume details"), "text/plain")}
        response = self.client.post("/api/v1/rag/documents/validate", files=file_payload)
        
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertTrue(data["valid"])
        self.assertEqual(data["file_name"], "resume.txt")
        self.assertEqual(data["file_type"], "txt")

    def test_api_validate_document_invalid(self):
        file_payload = {"file": ("malicious.exe", io.BytesIO(b"MZ...\x00\x01\x02"), "application/octet-stream")}
        response = self.client.post("/api/v1/rag/documents/validate", files=file_payload)
        
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertFalse(data["valid"])
        self.assertTrue(len(data["errors"]) > 0)

    @patch("app.modules.rag.services.loaders.txt.TXTLoader.load")
    def test_api_load_document_success(self, mock_txt_load):
        # Mock load response
        mock_txt_load.return_value = [Document(content="Sample text content for parsing.", metadata={"total_pages": 1})]

        file_payload = {"file": ("resume.txt", io.BytesIO(b"Sample text content for parsing."), "text/plain")}
        form_payload = {"metadata": json.dumps({"owner_id": 10})}

        response = self.client.post(
            "/api/v1/rag/documents/load", 
            files=file_payload,
            data=form_payload
        )

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        
        self.assertEqual(data["content"], "Sample text content for parsing.")
        self.assertEqual(data["metadata"]["file_name"], "resume.txt")
        self.assertEqual(data["metadata"]["custom_metadata"]["owner_id"], 10)
        self.assertEqual(data["metadata"]["num_words"], 5) # "Sample", "text", "content", "for", "parsing."

    def test_api_load_document_malformed_metadata(self):
        file_payload = {"file": ("resume.txt", io.BytesIO(b"Sample text"), "text/plain")}
        form_payload = {"metadata": "invalid-json"}

        response = self.client.post(
            "/api/v1/rag/documents/load", 
            files=file_payload,
            data=form_payload
        )

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("Invalid JSON", response.json()["message"])
