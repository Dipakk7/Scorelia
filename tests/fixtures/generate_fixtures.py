import os
import fitz
import docx

def generate_fixtures():
    fixtures_dir = os.path.dirname(os.path.abspath(__file__))
    os.makedirs(fixtures_dir, exist_ok=True)

    # 1. Generate sample_resume.pdf
    pdf_path = os.path.join(fixtures_dir, "sample_resume.pdf")
    doc = fitz.open()
    page = doc.new_page()
    text = (
        "John Doe\n"
        "Software Engineer\n"
        "john.doe@email.com | (123) 456-7890\n"
        "LinkedIn: https://linkedin.com/in/johndoe | GitHub: https://github.com/johndoe\n"
        "LeetCode: https://leetcode.com/johndoe | Website: https://johndoe.dev\n\n"
        "Skills:\n"
        "Python\n"
        "FastAPI\n"
        "SQL\n"
        "Docker\n"
        "PyTorch\n"
        "Git\n\n"
        "Education:\n"
        "B.Tech in Computer Science\n"
        "Stanford University, 2020\n"
        "Master of Science in Artificial Intelligence\n"
        "MIT, 2022\n\n"
        "Experience:\n"
        "Software Engineer at Google\n"
        "Jan 2020 - Present\n"
        "- Developed microservices using FastAPI and Python.\n"
        "- Optimized database performance using SQL.\n"
        "ML Engineer at Meta\n"
        "Jun 2019 - Aug 2019\n"
        "- Researched Deep Learning models using PyTorch.\n\n"
        "Projects:\n"
        "CareerPilot AI: Side project for parsing resumes.\n"
        "- Built parser backend with Python and FastAPI.\n\n"
        "Certifications:\n"
        "Oracle AI Foundations\n"
        "OCI Architect Associate"
    )
    # Write text to PDF
    page.insert_text((50, 50), text)
    doc.save(pdf_path)
    doc.close()
    print(f"Generated {pdf_path}")

    # 2. Generate sample_resume.docx
    docx_path = os.path.join(fixtures_dir, "sample_resume.docx")
    doc_docx = docx.Document()
    doc_docx.add_paragraph("John Doe")
    doc_docx.add_paragraph("Software Engineer")
    doc_docx.add_paragraph("john.doe@email.com | (123) 456-7890")
    doc_docx.add_paragraph("LinkedIn: https://linkedin.com/in/johndoe | GitHub: https://github.com/johndoe")
    doc_docx.add_paragraph("LeetCode: https://leetcode.com/johndoe | Website: https://johndoe.dev")
    
    doc_docx.add_heading("Skills", level=1)
    doc_docx.add_paragraph("Python")
    doc_docx.add_paragraph("FastAPI")
    doc_docx.add_paragraph("SQL")
    doc_docx.add_paragraph("Docker")
    doc_docx.add_paragraph("PyTorch")
    doc_docx.add_paragraph("Git")
    
    doc_docx.add_heading("Education", level=1)
    doc_docx.add_paragraph("B.Tech in Computer Science")
    doc_docx.add_paragraph("Stanford University, 2020")
    doc_docx.add_paragraph("Master of Science in Artificial Intelligence")
    doc_docx.add_paragraph("MIT, 2022")
    
    doc_docx.add_heading("Experience", level=1)
    doc_docx.add_paragraph("Software Engineer at Google")
    doc_docx.add_paragraph("Jan 2020 - Present")
    doc_docx.add_paragraph("- Developed microservices using FastAPI and Python.")
    doc_docx.add_paragraph("- Optimized database performance using SQL.")
    doc_docx.add_paragraph("ML Engineer at Meta")
    doc_docx.add_paragraph("Jun 2019 - Aug 2019")
    doc_docx.add_paragraph("- Researched Deep Learning models using PyTorch.")
    
    doc_docx.add_heading("Projects", level=1)
    doc_docx.add_paragraph("CareerPilot AI: Side project for parsing resumes.")
    doc_docx.add_paragraph("- Built parser backend with Python and FastAPI.")
    
    doc_docx.add_heading("Certifications", level=1)
    doc_docx.add_paragraph("Oracle AI Foundations")
    doc_docx.add_paragraph("OCI Architect Associate")
    doc_docx.save(docx_path)
    print(f"Generated {docx_path}")

    # 3. Generate corrupt_file.pdf
    corrupt_path = os.path.join(fixtures_dir, "corrupt_file.pdf")
    with open(corrupt_path, "wb") as f:
        f.write(b"This is a corrupt pdf file with junk bytes instead of real PDF content!")
    print(f"Generated {corrupt_path}")

    # 4. Generate empty_file.pdf
    empty_path = os.path.join(fixtures_dir, "empty_file.pdf")
    with open(empty_path, "wb") as f:
        pass
    print(f"Generated {empty_path}")

    # 5. Generate password_protected.pdf
    pw_path = os.path.join(fixtures_dir, "password_protected.pdf")
    doc_pw = fitz.open()
    page_pw = doc_pw.new_page()
    page_pw.insert_text((50, 50), "This is a password-protected resume.")
    doc_pw.save(pw_path, encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="password", owner_pw="owner")
    doc_pw.close()
    print(f"Generated {pw_path}")

    # 6. Generate scanned_resume.pdf (blank/no extractable text)
    scanned_path = os.path.join(fixtures_dir, "scanned_resume.pdf")
    doc_scanned = fitz.open()
    doc_scanned.new_page() # empty page with no text
    doc_scanned.save(scanned_path)
    doc_scanned.close()
    print(f"Generated {scanned_path}")

if __name__ == "__main__":
    generate_fixtures()
