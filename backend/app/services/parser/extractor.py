"""Text extraction modules for PDF and DOCX documents."""

import os
import fitz
import docx
import structlog
from app.core.config import settings
from app.models.resume import Resume
from app.services.parser.cleaner import clean_text

logger = structlog.get_logger()

def extract_pdf_text(file_path: str) -> str:
    """Extract raw text from a PDF file.

    Args:
        file_path: The absolute or relative path to the PDF file.

    Returns:
        The extracted raw text.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    if not os.path.isfile(file_path):
        raise ValueError(f"Path is not a file: {file_path}")

    try:
        with fitz.open(file_path) as document:
            logger.info("resume_pdf_loaded", file_path=file_path)

            if document.needs_pass:
                raise ValueError("Password protected PDF is not supported.")

            if len(document) > settings.MAX_PAGES:
                raise ValueError(f"PDF exceeds maximum allowed pages of {settings.MAX_PAGES}.")

            if len(document) == 0:
                raise ValueError("PDF has zero pages.")

            pages_text = []
            for page in document:
                text = page.get_text()
                if text:
                    pages_text.append(text)

            raw_text = "\n".join(pages_text)

            # Check if there is any extractable text (scanned image PDF)
            if not raw_text.strip():
                logger.warning("resume_pdf_no_text", file_path=file_path)
                return ""

            return raw_text
    except ValueError as ve:
        # Re-raise explicit validation exceptions
        raise ve
    except Exception as e:
        logger.error("resume_pdf_extraction_failed", file_path=file_path, error=str(e))
        raise ValueError("Unable to read PDF file.") from e

def extract_docx_text(file_path: str) -> str:
    """Extract raw text from a DOCX file.

    Args:
        file_path: The absolute or relative path to the DOCX file.

    Returns:
        The extracted raw text.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    if not os.path.isfile(file_path):
        raise ValueError(f"Path is not a file: {file_path}")

    try:
        doc = docx.Document(file_path)
        logger.info("resume_docx_loaded", file_path=file_path)

        text_parts = []

        # 1. Section Headers
        for section in doc.sections:
            header = section.header
            if header and not header.is_linked_to_previous:
                for p in header.paragraphs:
                    if p.text:
                        text_parts.append(p.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text:
                                    text_parts.append(p.text)

        # 2. Body Paragraphs
        for p in doc.paragraphs:
            if p.text:
                text_parts.append(p.text)

        # 3. Body Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            text_parts.append(p.text)

        # 4. Section Footers
        for section in doc.sections:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for p in footer.paragraphs:
                    if p.text:
                        text_parts.append(p.text)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text:
                                    text_parts.append(p.text)

        raw_text = "\n".join(text_parts)
        return raw_text
    except Exception as e:
        logger.error("resume_docx_extraction_failed", file_path=file_path, error=str(e))
        raise ValueError("Unable to read DOCX file.") from e

def extract_text(resume: Resume) -> str:
    """Extract and clean raw text from a Resume document based on its file type.

    Args:
        resume: The Resume database model instance.

    Returns:
        The cleaned extracted text string.
    """
    file_type = resume.file_type.lower().strip(".")
    if file_type == "pdf":
        raw_text = extract_pdf_text(resume.file_path)
    elif file_type == "docx":
        raw_text = extract_docx_text(resume.file_path)
    else:
        raise ValueError("Unsupported file type.")

    cleaned = clean_text(raw_text)
    return cleaned
