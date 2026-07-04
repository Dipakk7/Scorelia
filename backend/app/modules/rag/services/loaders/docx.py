# app/modules/rag/services/loaders/docx.py

import os
import docx
from typing import List, Generator, Dict, Any
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.modules.rag.services.loaders.base import BaseDocumentLoader
from app.modules.rag.schemas.document import Document
from app.modules.rag.exceptions import CorruptedFileError, InvalidDocumentError


class DOCXLoader(BaseDocumentLoader):
    """Document loader for Microsoft Word (.docx) files using python-docx."""

    def load(self, file_path: str, custom_metadata: Dict[str, Any] = None) -> List[Document]:
        """Completely load the DOCX and return a list of Documents."""
        return list(self.lazy_load(file_path, custom_metadata))

    def lazy_load(self, file_path: str, custom_metadata: Dict[str, Any] = None) -> Generator[Document, None, None]:
        """Lazily load the DOCX file, yielding segments."""
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise CorruptedFileError(f"Failed to open DOCX file. It might be corrupted: {str(e)}") from e

        # Extract metadata from core properties
        doc_properties = {}
        try:
            props = doc.core_properties
            if props.title:
                doc_properties["title"] = props.title
            if props.author:
                doc_properties["author"] = props.author
            if props.category:
                doc_properties["category"] = props.category
            if props.subject:
                doc_properties["subject"] = props.subject
        except Exception:
            pass

        paragraph_count = 0
        text_parts = []

        def format_paragraph(p: Paragraph) -> str:
            """Formats paragraph text preserving bold/italic formatting as markdown."""
            parts = []
            for run in p.runs:
                text = run.text
                if not text:
                    continue
                if run.bold and run.italic:
                    parts.append(f"***{text}***")
                elif run.bold:
                    parts.append(f"**{text}**")
                elif run.italic:
                    parts.append(f"*{text}*")
                else:
                    parts.append(text)
            return "".join(parts) if p.runs else p.text

        def iter_block_items(parent):
            """Iterates through paragraphs and tables in logical document order."""
            if isinstance(parent, DocxDocumentType):
                parent_elm = parent.element.body
            else:
                parent_elm = parent._element

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                paragraph_count += 1
                text = format_paragraph(block)
                if not text.strip():
                    continue

                # Translate heading styles to markdown header markup
                style_name = block.style.name if block.style else ""
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                        # Keep level within standard bounds (1-6)
                        level = max(1, min(6, level))
                        text = f"{'#' * level} {text}"
                    except Exception:
                        text = f"# {text}"

                text_parts.append(text)

            elif isinstance(block, Table):
                table_text = []
                for row in block.rows:
                    # Strip, and format cells
                    row_cells = []
                    for cell in row.cells:
                        # Extract paragraphs inside the cell and join them
                        cell_parts = []
                        for cell_p in cell.paragraphs:
                            cell_parts.append(format_paragraph(cell_p))
                        cell_text_clean = " ".join(cell_parts).strip().replace("\n", " ")
                        row_cells.append(cell_text_clean)
                    
                    table_text.append(" | ".join(row_cells))
                
                if table_text:
                    # Wrap table block with blank lines
                    text_parts.append("\n" + "\n".join(table_text) + "\n")

        full_text = "\n\n".join(text_parts)

        # Standard document page properties (DOCX is layout-flowed, so total_pages defaults to 1)
        meta = {
            "paragraph_count": paragraph_count,
            "total_pages": 1,
        }
        meta.update(doc_properties)
        if custom_metadata:
            meta.update(custom_metadata)

        yield Document(content=full_text, metadata=meta)
