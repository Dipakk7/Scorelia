import os
import uuid
import structlog
from datetime import datetime
from io import BytesIO
from typing import Optional, List, Dict, Any
from sqlalchemy.orm import Session

# reportlab (PDF)
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY

# python-docx (DOCX)
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.cover_letter.crud import crud
from app.cover_letter.models.ai_cover_letter import AICoverLetter
from app.cover_letter.models.ai_cover_letter_optimization import AICoverLetterOptimization
from app.cover_letter.models.ai_cover_letter_export import AICoverLetterExport
from app.models.resume import Resume
from app.models.user import User
from app.cover_letter.utils.templates import TEMPLATES

logger = structlog.get_logger()

def add_p_border_bottom(paragraph):
    """Draws a professional thin bottom border under a paragraph in DOCX."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '94A3B8')  # slate-400
    pBdr.append(bottom)
    pPr.append(pBdr)

class CoverLetterExportService:
    """Service layer managing the lifecycle and file rendering of AI cover letter exports."""

    def __init__(self, db: Session):
        self.db = db
        self.export_dir = "storage/cover_letter_exports"
        os.makedirs(self.export_dir, exist_ok=True)

    async def get_export(self, export_id: uuid.UUID, user_id: uuid.UUID) -> Optional[AICoverLetterExport]:
        """Fetch an export record by ID and check ownership."""
        return crud.get_export_by_id(self.db, export_id, user_id)

    async def get_exports(self, user_id: uuid.UUID) -> List[AICoverLetterExport]:
        """Retrieve export record list for the user."""
        return crud.get_exports_by_user_id(self.db, user_id)

    async def delete_export(self, export_id: uuid.UUID, user_id: uuid.UUID) -> bool:
        """Deletes export record and cleans up local storage file."""
        export_rec = crud.get_export_by_id(self.db, export_id, user_id)
        if not export_rec:
            return False

        # Physical file cleanup
        file_path = os.path.join(self.export_dir, f"{export_rec.id}.{export_rec.export_format.lower()}")
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info("deleted_physical_export_file", file_path=file_path)
            except Exception as e:
                logger.error("failed_to_delete_physical_export_file", file_path=file_path, error=str(e))

        return crud.delete_export_record(self.db, export_rec)

    async def export_cover_letter(
        self,
        user_id: uuid.UUID,
        cover_letter_id: uuid.UUID,
        export_format: str,
        template_name: str = "Professional",
        optimization_id: Optional[uuid.UUID] = None
    ) -> AICoverLetterExport:
        """
        Coordinates full validation, content selection, template styling, file generation,
        saving to disk, and database tracking.
        """
        # 1. Supported formats validation
        format_upper = export_format.upper().strip()
        if format_upper not in ["PDF", "DOCX", "MD", "TXT"]:
            raise ValueError(f"Unsupported export format: {export_format}")

        # 2. Retrieve and validate cover letter
        cover_letter = crud.get_cover_letter_by_id(self.db, cover_letter_id, user_id)
        if not cover_letter:
            raise ValueError("Cover letter not found or not owned by user.")

        # 3. Retrieve and validate optimization (if provided)
        optimization = None
        if optimization_id:
            optimization = crud.get_optimization_by_id(self.db, optimization_id, user_id)
            if not optimization:
                raise ValueError("Optimization record not found or not owned by user.")
            if optimization.cover_letter_id != cover_letter_id:
                raise ValueError("Optimization record does not correspond to the requested cover letter.")

        # 4. Resolve template style
        template_key = template_name.strip().title()
        if template_key not in TEMPLATES:
            raise ValueError(f"Unsupported template name: {template_name}")
        template_style = TEMPLATES[template_key]

        # 5. Determine correct content (optimized vs original)
        if optimization:
            content_text = optimization.optimization_result.get("optimized_content")
        else:
            content_text = cover_letter.generated_content

        if not content_text or not content_text.strip():
            raise ValueError("Cannot export a cover letter with empty content.")

        # 6. Parse cover letter sections
        parsed_sections = self._parse_cover_letter_content(content_text)

        # 7. Extract contact details
        contact_info = self._get_contact_info(cover_letter)

        # 8. Date and filenames
        date_str = datetime.utcnow().strftime("%B %d, %Y")
        clean_company = "".join(c for c in cover_letter.company_name if c.isalnum() or c in (" ", "_", "-")).strip().replace(" ", "_")
        friendly_filename = f"cover_letter_{clean_company}_{template_key.lower()}.{format_upper.lower()}"

        # 9. Format metadata properties
        meta_dict = {
            "export_time": datetime.utcnow().isoformat() + "Z",
            "export_version": "1.0.0",
            "writing_style": cover_letter.writing_style,
            "generation_mode": cover_letter.generation_mode,
            "optimization_score": optimization.quality_score if optimization else None,
            "provider": optimization.provider if optimization else cover_letter.provider,
            "model": optimization.model if optimization else cover_letter.model
        }

        # Build clean string representation for comments
        opt_score_part = f" | Score: {optimization.quality_score}" if optimization else ""
        model_part = meta_dict["model"] if meta_dict["model"] else "Unknown"
        provider_part = meta_dict["provider"] if meta_dict["provider"] else "Unknown"
        metadata_comment = (
            f"Export Time: {meta_dict['export_time']} | "
            f"Style: {meta_dict['writing_style']} | "
            f"Mode: {meta_dict['generation_mode']}{opt_score_part} | "
            f"LLM: {model_part} ({provider_part})"
        )

        # 10. Generate file binary contents
        if format_upper == "PDF":
            file_data = self._generate_pdf(parsed_sections, contact_info, cover_letter.company_name, cover_letter.job_title, template_style, date_str, metadata_comment)
        elif format_upper == "DOCX":
            file_data = self._generate_docx(parsed_sections, contact_info, cover_letter.company_name, cover_letter.job_title, template_style, date_str, metadata_comment)
        elif format_upper == "MD":
            file_data = self._generate_markdown(parsed_sections, contact_info, cover_letter.company_name, cover_letter.job_title, date_str, metadata_comment)
        else:  # TXT
            file_data = self._generate_text(parsed_sections, contact_info, cover_letter.company_name, cover_letter.job_title, date_str, metadata_comment)

        # 11. Write physically to storage using export record UUID placeholder
        export_id = uuid.uuid4()
        storage_filename = f"{export_id}.{format_upper.lower()}"
        file_path = os.path.join(self.export_dir, storage_filename)
        
        try:
            with open(file_path, "wb") as f:
                f.write(file_data)
        except Exception as e:
            logger.error("failed_to_write_export_file_to_disk", file_path=file_path, error=str(e))
            raise RuntimeError(f"Failed to write export file to disk: {str(e)}")

        # 12. Create database tracking record
        db_export = AICoverLetterExport(
            id=export_id,
            user_id=user_id,
            cover_letter_id=cover_letter_id,
            optimization_id=optimization_id,
            export_format=format_upper,
            template_name=template_key,
            file_name=friendly_filename,
            file_size=len(file_data),
            export_metadata=meta_dict,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        self.db.add(db_export)
        self.db.commit()
        self.db.refresh(db_export)

        logger.info(
            "cover_letter_exported_successfully",
            export_id=str(db_export.id),
            format=format_upper,
            template=template_key,
            file_size=len(file_data)
        )

        return db_export

    def _parse_cover_letter_content(self, content: str) -> dict:
        """Robust parser dividing unstructured or structured cover letter text blocks into formatting layers."""
        paragraphs = [p.strip() for p in content.replace("\r\n", "\n").split("\n\n") if p.strip()]
        result = {
            "title": None,
            "greeting": "Dear Hiring Team,",
            "body": [],
            "closing": "Sincerely,",
            "signature": "Applicant"
        }

        if not paragraphs:
            return result

        idx = 0
        first = paragraphs[idx]

        # 1. Subject/Title (Must not start with Dear, be relatively short)
        if first.lower().startswith("subject:") or first.lower().startswith("application for") or (len(first) < 70 and not first.lower().startswith("dear")):
            result["title"] = first
            idx += 1

        # 2. Greeting check
        if idx < len(paragraphs):
            greet_cand = paragraphs[idx]
            if greet_cand.lower().startswith("dear") or greet_cand.lower().startswith("to ") or greet_cand.endswith(",") or greet_cand.endswith(":"):
                result["greeting"] = greet_cand
                idx += 1

        remaining = paragraphs[idx:]

        # 3. Separating Body, Closing, and Signature from the tail
        if len(remaining) >= 2:
            last_two = remaining[-2:]
            closing_words = ["sincerely", "best", "regards", "respectfully", "thank", "yours", "appreciate"]
            looks_like_closing = any(w in last_two[0].lower() for w in closing_words) or len(last_two[0]) < 25
            if looks_like_closing:
                result["closing"] = last_two[0]
                result["signature"] = last_two[1]
                result["body"] = remaining[:-2]
            else:
                result["body"] = remaining
        elif len(remaining) == 1:
            result["body"] = remaining
        
        return result

    def _get_contact_info(self, cover_letter: AICoverLetter) -> dict:
        """Extract user contact info from resume parsed_data or user account details."""
        resume = self.db.query(Resume).filter(Resume.id == cover_letter.resume_id).first()
        user = self.db.query(User).filter(User.id == cover_letter.user_id).first()

        info = {
            "name": user.full_name if user else "Applicant",
            "email": user.email if user else "",
            "phone": "",
            "links": []
        }

        if resume and resume.parsed_data:
            personal = resume.parsed_data.get("personal_info", {})
            if personal:
                if personal.get("name"):
                    info["name"] = personal.get("name")
                if personal.get("email"):
                    info["email"] = personal.get("email")
                if personal.get("phone"):
                    info["phone"] = personal.get("phone")
                
                # Retrieve links
                links = personal.get("links") or personal.get("websites") or []
                if isinstance(links, list):
                    info["links"] = [str(lnk) for lnk in links]
                elif isinstance(links, str):
                    info["links"] = [links]

        return info

    def _generate_pdf(self, content_parts: dict, contact: dict, company_name: str, job_title: str, template: dict, date_str: str, metadata_comment: str) -> bytes:
        """Renders cover letter to a professionally styled PDF byte stream using ReportLab."""
        buffer = BytesIO()
        margin = template["margin_inches"] * 72  # Convert to points

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=margin,
            rightMargin=margin,
            topMargin=margin,
            bottomMargin=margin
        )

        def add_pdf_metadata(canvas, doc):
            canvas.setAuthor(contact["name"])
            canvas.setTitle(f"Cover Letter - {company_name}")
            canvas.setSubject(f"Application for {job_title}")
            canvas.setCreator("CareerPilot AI")
            canvas.setKeywords(metadata_comment)

        story = []
        styles = getSampleStyleSheet()

        font_family = template["font_family"]
        if font_family == "Times-Roman":
            font_bold = "Times-Bold"
        else:
            font_bold = f"{font_family}-Bold"
        
        primary_color = colors.HexColor(template["primary_color"])
        secondary_color = colors.HexColor(template["secondary_color"])

        # Styled components
        name_style = ParagraphStyle(
            'HeaderName',
            parent=styles['Normal'],
            fontName=font_bold,
            fontSize=22,
            leading=26,
            textColor=primary_color,
            spaceAfter=4,
            alignment=TA_LEFT if template["header_align"] == "LEFT" else (TA_RIGHT if template["header_align"] == "RIGHT" else TA_CENTER)
        )

        contact_parts = []
        if contact["email"]: contact_parts.append(contact["email"])
        if contact["phone"]: contact_parts.append(contact["phone"])
        contact_parts.extend(contact["links"][:2])
        contact_text = " | ".join(contact_parts)

        contact_style = ParagraphStyle(
            'HeaderContact',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=9,
            leading=12,
            textColor=secondary_color,
            spaceAfter=12,
            alignment=name_style.alignment
        )

        # Header Block
        story.append(Paragraph(contact["name"], name_style))
        if contact_text:
            story.append(Paragraph(contact_text, contact_style))

        # Horizontal Divider line
        if template["divider"]:
            story.append(HRFlowable(width="100%", thickness=1, color=secondary_color, spaceAfter=template["spacing_after_header"], spaceBefore=0))
        else:
            story.append(Spacer(1, template["spacing_after_header"]))

        # Recipient Info
        recipient_style = ParagraphStyle(
            'RecipientStyle',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=template["font_size"],
            leading=template["line_height"],
            textColor=colors.HexColor("#1E293B"),
            spaceAfter=12
        )
        recipient_text = f"<b>Date:</b> {date_str}<br/><b>To:</b> Hiring Team / Recruitment Group<br/><b>Company:</b> {company_name}<br/><b>Position:</b> {job_title}"
        story.append(Paragraph(recipient_text, recipient_style))

        # Subject Line / Title
        if content_parts.get("title"):
            title_style = ParagraphStyle(
                'CLTitle',
                parent=styles['Normal'],
                fontName=font_bold,
                fontSize=template["font_size"] + 1,
                leading=template["line_height"] + 2,
                textColor=primary_color,
                spaceAfter=14
            )
            story.append(Paragraph(content_parts["title"], title_style))

        # Greeting
        greeting_style = ParagraphStyle(
            'CLGreeting',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=template["font_size"],
            leading=template["line_height"],
            textColor=colors.HexColor("#000000"),
            spaceAfter=10
        )
        story.append(Paragraph(content_parts["greeting"], greeting_style))

        # Body Paragraphs
        body_style = ParagraphStyle(
            'CLBody',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=template["font_size"],
            leading=template["line_height"],
            textColor=colors.HexColor("#111827"),
            spaceAfter=template["spacing_paragraph"],
            alignment=TA_JUSTIFY if font_family == "Times-Roman" else TA_LEFT
        )

        for para in content_parts["body"]:
            story.append(Paragraph(para.replace("\n", "<br/>"), body_style))

        # Closing / Sign-off
        closing_style = ParagraphStyle(
            'CLClosing',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=template["font_size"],
            leading=template["line_height"],
            textColor=colors.HexColor("#000000"),
            spaceAfter=12
        )

        signature_style = ParagraphStyle(
            'CLSignature',
            parent=styles['Normal'],
            fontName=font_bold,
            fontSize=template["font_size"],
            leading=template["line_height"],
            textColor=colors.HexColor("#000000")
        )

        signature_block = [
            Paragraph(content_parts["closing"], closing_style),
            Spacer(1, 15),
            Paragraph(content_parts["signature"], signature_style)
        ]
        story.append(KeepTogether(signature_block))

        # Footer Notice
        story.append(Spacer(1, 20))
        footer_style = ParagraphStyle(
            'CLFooter',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#94A3B8"),
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Generated with CareerPilot AI | {metadata_comment}", footer_style))

        doc.build(story, onFirstPage=add_pdf_metadata)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def _generate_docx(self, content_parts: dict, contact: dict, company_name: str, job_title: str, template: dict, date_str: str, metadata_comment: str) -> bytes:
        """Renders cover letter to a professionally styled Word document byte stream using python-docx."""
        doc = docx.Document()

        # Page margins
        margin = template["margin_inches"]
        for section in doc.sections:
            section.top_margin = Inches(margin)
            section.bottom_margin = Inches(margin)
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)

        def hex_to_rgb(hex_str: str) -> RGBColor:
            h = hex_str.lstrip('#')
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        primary_color = hex_to_rgb(template["primary_color"])
        secondary_color = hex_to_rgb(template["secondary_color"])
        font_family = template["font_family"]

        # Title name header
        p_name = doc.add_paragraph()
        if template["header_align"] == "CENTER":
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif template["header_align"] == "RIGHT":
            p_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run_name = p_name.add_run(contact["name"])
        run_name.font.name = font_family
        run_name.font.size = Pt(22)
        run_name.font.bold = True
        run_name.font.color.rgb = primary_color
        p_name.paragraph_format.space_after = Pt(2)

        # Contact line
        contact_parts = []
        if contact["email"]: contact_parts.append(contact["email"])
        if contact["phone"]: contact_parts.append(contact["phone"])
        contact_parts.extend(contact["links"][:2])
        contact_text = " | ".join(contact_parts)

        p_contact = None
        if contact_text:
            p_contact = doc.add_paragraph()
            p_contact.alignment = p_name.alignment
            run_contact = p_contact.add_run(contact_text)
            run_contact.font.name = font_family
            run_contact.font.size = Pt(9)
            run_contact.font.color.rgb = secondary_color
            p_contact.paragraph_format.space_after = Pt(12)

        # Bottom Border
        if template["divider"]:
            add_p_border_bottom(p_contact if p_contact else p_name)

        # Spacer after header
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(template["spacing_after_header"])
        p_space.paragraph_format.space_after = Pt(0)

        # Recipient Info
        p_rec = doc.add_paragraph()
        p_rec.paragraph_format.line_spacing = 1.15
        p_rec.paragraph_format.space_after = Pt(14)
        run_rec = p_rec.add_run(f"Date: {date_str}\nTo: Hiring Team / Recruitment Group\nCompany: {company_name}\nPosition: {job_title}")
        run_rec.font.name = font_family
        run_rec.font.size = Pt(template["font_size"])
        run_rec.font.color.rgb = hex_to_rgb("#1E293B")

        # Subject Title
        if content_parts.get("title"):
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_after = Pt(12)
            run_title = p_title.add_run(content_parts["title"])
            run_title.font.name = font_family
            run_title.font.size = Pt(template["font_size"] + 1)
            run_title.font.bold = True
            run_title.font.color.rgb = primary_color

        # Greeting
        p_greet = doc.add_paragraph()
        p_greet.paragraph_format.space_after = Pt(10)
        run_greet = p_greet.add_run(content_parts["greeting"])
        run_greet.font.name = font_family
        run_greet.font.size = Pt(template["font_size"])

        # Body paragraphs
        for para in content_parts["body"]:
            p_body = doc.add_paragraph()
            p_body.paragraph_format.space_after = Pt(template["spacing_paragraph"])
            p_body.paragraph_format.line_spacing = float(template["line_height"]) / template["font_size"]
            if font_family == "Times-Roman":
                p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run_body = p_body.add_run(para)
            run_body.font.name = font_family
            run_body.font.size = Pt(template["font_size"])

        # Closing
        p_close = doc.add_paragraph()
        p_close.paragraph_format.space_after = Pt(15)
        run_close = p_close.add_run(content_parts["closing"])
        run_close.font.name = font_family
        run_close.font.size = Pt(template["font_size"])

        # Signature
        p_sig = doc.add_paragraph()
        p_sig.paragraph_format.space_after = Pt(30)
        run_sig = p_sig.add_run(content_parts["signature"])
        run_sig.font.name = font_family
        run_sig.font.size = Pt(template["font_size"])
        run_sig.font.bold = True

        # Footer Notice
        p_foot = doc.add_paragraph()
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_foot = p_foot.add_run(f"Generated with CareerPilot AI | {metadata_comment}")
        run_foot.font.name = font_family
        run_foot.font.size = Pt(7)
        run_foot.font.color.rgb = hex_to_rgb("#94A3B8")

        # Document core metadata
        doc.core_properties.author = contact["name"]
        doc.core_properties.title = f"Cover Letter - {company_name}"
        doc.core_properties.comments = metadata_comment

        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    def _generate_markdown(self, content_parts: dict, contact: dict, company_name: str, job_title: str, date_str: str, metadata_comment: str) -> bytes:
        """Renders cover letter to standard Markdown bytes."""
        contact_parts = []
        if contact["email"]: contact_parts.append(contact["email"])
        if contact["phone"]: contact_parts.append(contact["phone"])
        contact_parts.extend(contact["links"][:2])
        contact_text = " | ".join(contact_parts)

        md_text = f"# {contact['name']}\n"
        if contact_text:
            md_text += f"{contact_text}\n"
        md_text += "\n---\n\n"
        md_text += f"**Date:** {date_str}  \n"
        md_text += "**To:** Hiring Team / Recruitment Group  \n"
        md_text += f"**Company:** {company_name}  \n"
        md_text += f"**Position:** {job_title}  \n\n"

        if content_parts.get("title"):
            md_text += f"### {content_parts['title']}\n\n"

        md_text += f"{content_parts['greeting']}\n\n"

        for para in content_parts["body"]:
            md_text += f"{para}\n\n"

        md_text += f"{content_parts['closing']}\n\n"
        md_text += f"**{content_parts['signature']}**\n\n"
        md_text += f"---\n*Generated with CareerPilot AI | {metadata_comment}*\n"

        return md_text.encode('utf-8')

    def _generate_text(self, content_parts: dict, contact: dict, company_name: str, job_title: str, date_str: str, metadata_comment: str) -> bytes:
        """Renders cover letter to plain UTF-8 text bytes."""
        contact_parts = []
        if contact["email"]: contact_parts.append(contact["email"])
        if contact["phone"]: contact_parts.append(contact["phone"])
        contact_parts.extend(contact["links"][:2])
        contact_text = " | ".join(contact_parts)

        txt = f"{contact['name']}\n"
        if contact_text:
            txt += f"{contact_text}\n"
        txt += "=" * 50 + "\n\n"
        txt += f"Date: {date_str}\n"
        txt += "To: Hiring Team / Recruitment Group\n"
        txt += f"Company: {company_name}\n"
        txt += f"Position: {job_title}\n\n"

        if content_parts.get("title"):
            txt += f"{content_parts['title']}\n\n"

        txt += f"{content_parts['greeting']}\n\n"

        for para in content_parts["body"]:
            txt += f"{para}\n\n"

        txt += f"{content_parts['closing']}\n\n"
        txt += f"{content_parts['signature']}\n\n"
        txt += "=" * 50 + "\n"
        txt += f"Generated with CareerPilot AI | {metadata_comment}\n"

        return txt.encode('utf-8')
